#!/usr/bin/env python3
"""
redact.py

PII Redaction Tool for the Red Herring Prospectus (or any .docx).

Usage:
    python3 redact.py <input.docx> <output.docx> [--audit audit_log.csv]

Approach
--------
1. Walk every paragraph in the document body AND every table cell
   (including nested tables), de-duplicating by the underlying XML
   element's identity so that merged/vertically-spanned cells are only
   visited once (a naive walk revisits the same XML element multiple
   times for merged cells, which caused a double-redaction bug during
   development -- see BUGFIX_NOTES.md).
2. For each paragraph, run all PII detectors (pii_detectors.detect_all)
   against the paragraph's full text.
3. Replace detected spans with consistent fake values (pii_faker), then
   write the redacted text back into the paragraph's runs.
4. Log every redaction (original value, fake value, label, location) to
   a CSV audit trail.

Run-boundary caveat: Word splits paragraph text across multiple <w:r>
runs. This script rewrites the paragraph's runs by clearing them and
writing the fully-redacted text into the first run (copying that run's
formatting), which loses intra-paragraph formatting variation (e.g. if
only one word was bold) but guarantees correct redaction regardless of
where run boundaries fall relative to detected PII spans. Paragraph-level
formatting (style, alignment) and document-level formatting are untouched.

Manual line-break handling: python-docx represents a manual line break
(Word's Shift+Enter, XML <w:br/>) as a literal "\n" character when you
read paragraph.text. This script splits the redacted text on "\n" and
rebuilds real <w:br/> elements between segments (rather than writing a
flattened string into one run), so any paragraph that does contain a
manual break keeps it intact after redaction. Checked directly against
this document's XML: none of its paragraphs actually use <w:br/> (text
that looks line-broken, like the cover page's company-name/CIN heading,
is just ordinary word-wrap reacting to the container width) -- so this
code path is defensive/future-proofing for other documents rather than a
fix for an observed bug in this specific one.
"""

import sys
import csv
import argparse
from docx import Document
from docx.oxml.ns import qn

from pii_detectors import detect_all
from pii_faker import FakeValueGenerator


def iter_unique_paragraphs(doc):
    """Yield every paragraph in the document (body + all tables, including
    nested tables), de-duplicated by a stable XML tree-path identifier so
    merged table cells are only processed once.

    NOTE: an earlier version of this function deduplicated using
    id(paragraph._p) (the Python memory address of the lxml element). That
    is unsafe: python-docx/lxml create fresh proxy objects on each
    `.paragraphs` access, and once the previous proxy is garbage collected,
    CPython can recycle its memory address for a *different* element. In
    practice this caused completely unrelated paragraphs (e.g. a
    Registered Office address and, two cells later, an email/phone cell)
    to collide on the same id() and get wrongly skipped as "already
    visited" -- silently dropping real PII from redaction on the
    document's own cover-page table. Using the element's absolute XML
    tree path (stable, string-based, unaffected by object lifetime) fixes
    this while still correctly deduplicating genuinely merged cells (which
    share the same underlying XML element and therefore the same path).
    """
    seen = set()

    def path_of(p):
        return p._p.getroottree().getpath(p._p)

    def walk_block_items(parent):
        # parent can be the document body or a table cell
        if hasattr(parent, "paragraphs"):
            for p in parent.paragraphs:
                key = path_of(p)
                if key not in seen:
                    seen.add(key)
                    yield p
        if hasattr(parent, "tables"):
            for t in parent.tables:
                for row in t.rows:
                    for cell in row.cells:
                        yield from walk_block_items(cell)

    yield from walk_block_items(doc)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                yield from walk_block_items(cell)


def _copy_run_format(src_run, dst_run):
    """Copy the visual formatting we care about from src_run to dst_run so
    a newly-created run (from splitting on a line break) doesn't silently
    revert to default formatting."""
    dst_run.bold = src_run.bold
    dst_run.italic = src_run.italic
    dst_run.underline = src_run.underline
    try:
        dst_run.font.size = src_run.font.size
        dst_run.font.name = src_run.font.name
        dst_run.font.color.rgb = src_run.font.color.rgb
    except Exception:
        # Some runs don't have an explicit color / size set -- fine to skip.
        pass


def redact_paragraph(paragraph, faker, audit_rows, location):
    text = paragraph.text
    if not text.strip():
        return

    spans = detect_all(text)
    if not spans:
        return

    # Build the redacted string, right to left so earlier offsets stay valid
    new_text = text
    for start, end, label, original in sorted(spans, key=lambda s: -s[0]):
        fake = faker.get(label, original)
        new_text = new_text[:start] + fake + new_text[end:]
        audit_rows.append({
            "location": location,
            "label": label,
            "original": original,
            "fake": fake,
        })

    if new_text == text:
        return

    runs = paragraph.runs
    if not runs:
        # Paragraph has no runs (rare) -- add one
        paragraph.add_run(new_text)
        return

    # Rewrite runs: preserve the formatting of the first run, clear the
    # rest. If the redacted text still contains manual line breaks ("\n",
    # from an original <w:br/>), rebuild real <w:br/> elements between
    # segments instead of flattening them into plain text (see module
    # docstring above).
    first_run = runs[0]
    segments = new_text.split("\n")

    first_run.text = segments[0]
    for r in runs[1:]:
        r.text = ""

    anchor_run = first_run
    for seg in segments[1:]:
        anchor_run.add_break()  # inserts a real <w:br/>
        new_run = paragraph.add_run(seg)
        _copy_run_format(first_run, new_run)
        anchor_run = new_run  # so a third+ line break chains correctly


def redact_document(input_path, output_path, audit_path):
    doc = Document(input_path)
    faker = FakeValueGenerator()
    audit_rows = []

    count = 0
    for para in iter_unique_paragraphs(doc):
        count += 1
        redact_paragraph(para, faker, audit_rows, location=f"para#{count}")

    doc.save(output_path)

    with open(audit_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=["location", "label", "original", "fake"])
        writer.writeheader()
        writer.writerows(audit_rows)

    # Summary counts by label
    from collections import Counter
    counts = Counter(r["label"] for r in audit_rows)
    print(f"Processed {count} unique paragraphs.")
    print(f"Total redactions: {len(audit_rows)}")
    for label, c in sorted(counts.items(), key=lambda x: -x[1]):
        print(f"  {label}: {c}")
    print(f"Redacted document written to: {output_path}")
    print(f"Audit log written to: {audit_path}")


if __name__ == "__main__":
    ap = argparse.ArgumentParser(description="Redact PII from a .docx file.")
    ap.add_argument("input", help="Path to input .docx")
    ap.add_argument("output", help="Path to write redacted .docx")
    ap.add_argument("--audit", default="audit_log.csv", help="Path to write CSV audit log")
    args = ap.parse_args()

    redact_document(args.input, args.output, args.audit)
